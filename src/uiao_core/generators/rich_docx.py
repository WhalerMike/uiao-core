"""Rich DOCX generator for UIAO leadership briefing.

Migrated from scripts/generate_rich_docx.py into the uiao_core package.
Produces publication-quality DOCX with native Word styles, auto-TOC,
headers/footers with classification markings, embedded images, and
properly formatted compliance tables.

References: ADR-0004
"""

# isort: skip_file
from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

try:
    from docxtpl import DocxTemplate

    HAS_DOCXTPL = True
except ImportError:
    HAS_DOCXTPL = False

from uiao_core.utils.context import get_settings, load_context


logger = logging.getLogger(__name__)

_DEFAULT_IMAGE_WIDTH: Any = Inches(5.5)  # module-level singleton for ruff B008
# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


def _add_classification_header(doc: Document, classification: str = "CUI") -> None:
    """Add header/footer with classification marking to all sections."""
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.text = classification
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
        run.font.bold = True

        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.text = f"{classification} | UIAO Program | Generated {datetime.now():%Y-%m-%d}"
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.runs[0]
        fr.font.size = Pt(7)
        fr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def _add_heading(doc: Document, text: str, level: int = 1):
    """Add a heading with consistent navy styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    return h


def _add_narrative(doc: Document, text: Any) -> None:
    """Add narrative paragraph(s) with proper formatting."""
    if not text:
        return
    for para_text in str(text).split("\n\n"):
        p = doc.add_paragraph(para_text.strip())
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(11)


def _add_image_safe(doc: Document, image_name: str, visuals_dir: Path, width: Any = _DEFAULT_IMAGE_WIDTH) -> bool:
    """Add an image if it exists, skip gracefully if not."""
    img_path = visuals_dir / image_name
    if img_path.exists():
        doc.add_picture(str(img_path), width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    logger.warning("Image not found: %s", img_path)
    return False


def _add_compliance_table(doc: Document, matrix: list[dict]) -> None:
    """Add the Unified Compliance Matrix as a formatted Word table."""
    headers = [
        "UIAO Pillar",
        "CISA ZT Pillar",
        "Target Maturity",
        "NIST 800-53 Controls",
        "Mission Impact",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Medium Shading 1 Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    for entry in matrix:
        if not isinstance(entry, dict):
            continue
        row = table.add_row()
        row.cells[0].text = entry.get("pillar", "")
        row.cells[3].text = ", ".join(controls) if isinstance(controls, list) else str(controls)
        row.cells[4].text = entry.get("impact_statement", "")
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)


def _add_evidence_table(doc: Document) -> None:
    """Add the FedRAMP Audit Evidence Summary table."""
    evidence_map = [
        ("V1: Identity-to-IP Mapping", "U + A (The Gate)", "IA-2, AC-19, CM-8"),
        ("V2: INR Fabric", "O (The Network)", "AC-4"),
        ("V3: 20x Governance Loop", "Governance (The Hub)", "CA-7, IR-4"),
        ("V4: Modernization Atlas", "Strategy (The Journey)", "Program Vision / TIC 3.0"),
        ("V5: Cryptographic Trust Chain", "Security (The Lock)", "SC-8"),
    ]
    headers = ["Visual Title", "Architectural Pillar", "NIST Control(s)"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light List Accent 1"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    for title, pillar, controls in evidence_map:
        row = table.add_row()
        row.cells[0].text = title
        row.cells[1].text = pillar
        row.cells[2].text = controls
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)


# ---------------------------------------------------------------------------
# Scratch builder
# ---------------------------------------------------------------------------


def _build_from_scratch(context: dict, visuals_dir: Path) -> Document:
    """Build a complete styled DOCX programmatically."""
    doc = Document()
    lb = context.get("leadership_briefing", {})
    if not isinstance(lb, dict):
        lb = {}
    classification = context.get("classification", "CUI")
    _add_classification_header(doc, classification)

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Unified Identity-Addressing-Overlay Architecture (UIAO)")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run(f"Leadership Briefing - Version {context.get('version', '1.0')}")
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_p.add_run(f"Generated: {datetime.now():%B %d, %Y}")
    dr.font.size = Pt(10)
    dr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    doc.add_page_break()

    # TOC placeholder
    _add_heading(doc, "Table of Contents", level=1)
    toc_p = doc.add_paragraph()
    toc_p.add_run("[Table of Contents - right-click and select 'Update Field' in Word]")
    toc_p.runs[0].font.size = Pt(9)
    toc_p.runs[0].font.italic = True
    doc.add_page_break()

    # Sections
    _add_heading(doc, "Executive Summary")
    _add_narrative(doc, lb.get("executive_summary"))

    _add_heading(doc, "Program Overview")
    _add_narrative(doc, lb.get("program_overview"))

    _add_heading(doc, "Why Modernization Is Required")
    _add_narrative(doc, lb.get("modernization_need"))

    _add_heading(doc, "Program Vision")
    _add_narrative(doc, lb.get("program_vision"))

    # Five Control Planes
    _add_heading(doc, "The Five Control Planes")
    for i, plane in enumerate(lb.get("control_planes", []), 1):
        if isinstance(plane, dict):
            _add_heading(doc, f"{i}. {plane.get('name', '')}", level=2)
            _add_narrative(doc, plane.get("narrative"))

    # Seven Core Concepts
    _add_heading(doc, "Seven Core Concepts")
    for i, concept in enumerate(lb.get("core_concepts", []), 1):
        if isinstance(concept, dict):
            _add_heading(doc, f"{i}. {concept.get('name', '')}", level=2)
            _add_narrative(doc, concept.get("narrative"))

    _add_heading(doc, "Frozen State Analysis")
    _add_narrative(doc, lb.get("frozen_state"))

    _add_heading(doc, "Program Outcomes")
    _add_narrative(doc, lb.get("outcomes"))

    # Vibrant Visualizations
    doc.add_page_break()
    _add_heading(doc, "Vibrant Visualizations")
    visuals = [
        ("Modernization Journey", "uiao-vibrant-modernization-atlas.png"),
        ("FedRAMP 20x Governance Loop", "uiao-vibrant-20x-governance-hub.png"),
        ("Identity-to-IP Architecture", "uiao-vibrant-u-plus-a-mapping.png"),
    ]
    for title_text, img in visuals:
        _add_heading(doc, title_text, level=2)
        _add_image_safe(doc, img, visuals_dir)

    # Maturity radar chart
    radar_path = visuals_dir / "dynamic-maturity-radar.png"
    if radar_path.exists():
        doc.add_page_break()
        _add_heading(doc, "CISA Zero Trust Maturity Assessment")
        _add_image_safe(doc, "dynamic-maturity-radar.png", visuals_dir)

    # Mermaid-rendered architecture diagrams
    mermaid_dir = visuals_dir.parent / "mermaid"
    if mermaid_dir.is_dir():
        mermaid_pngs = sorted(mermaid_dir.glob("*.png"))
        if mermaid_pngs:
            doc.add_page_break()
            _add_heading(doc, "Architecture Diagrams (Mermaid)")
            for png in mermaid_pngs:
                _add_heading(doc, png.stem.replace("-", " ").title(), level=2)
                _add_image_safe(doc, png.name, mermaid_dir)

    # Gemini AI-generated visuals
    gemini_dir = visuals_dir.parent / "gemini"
    if gemini_dir.is_dir():
        gemini_pngs = sorted(gemini_dir.glob("*.png"))
        if gemini_pngs:
            doc.add_page_break()
            _add_heading(doc, "AI-Generated Visuals (Gemini)")
            for png in gemini_pngs:
                _add_heading(doc, png.stem.replace("-", " ").title(), level=2)
                _add_image_safe(doc, png.name, gemini_dir)

    # FedRAMP Evidence Summary
    doc.add_page_break()
    _add_heading(doc, "FedRAMP 20x Audit Evidence Summary")
    _add_narrative(
        doc,
        "Direct mapping of UIAO architecture to NIST 800-53 Rev 5 controls.",
    )
    _add_evidence_table(doc)

    # Compliance Matrix
    doc.add_page_break()
    _add_heading(doc, "Unified Compliance & Maturity Matrix")
    matrix = context.get("unified_compliance_matrix", [])
    if matrix:
        _add_compliance_table(doc, matrix)
        p = doc.add_paragraph()
        p.add_run(
            "Auditor Note: All controls listed above are continuously "
            "monitored via the UIAO Governance Plane (V3) and reported "
            "through the ServiceNow SCuBA integration."
        ).font.italic = True

    return doc


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def build_rich_docx(
    *,
    canon_path: Path | None = None,
    data_dir: Path | None = None,
    templates_dir: Path | None = None,
    visuals_dir: Path | None = None,
    exports_dir: Path | None = None,
) -> Path:
    """Generate a rich DOCX leadership briefing and return the output path.

    All directory arguments default to values derived from ``Settings``.
    """
    settings = get_settings()
    canon_path = canon_path or settings.canon_dir / "uiao_leadership_briefing_v1.0.yaml"
    data_dir = data_dir or settings.data_dir
    templates_dir = templates_dir or settings.templates_dir
    visuals_dir = visuals_dir or settings.visuals_dir
    exports_dir = exports_dir or settings.exports_dir

    logger.info("Loading UIAO context...")
    context = load_context(canon_path=canon_path, data_dir=data_dir)

    docx_dir = exports_dir / "docx"
    docx_dir.mkdir(parents=True, exist_ok=True)
    out_path = docx_dir / "UIAO_Leadership_Briefing_v1.0.docx"

    # Try template-based approach first
    tpl_path = templates_dir / "leadership_briefing_v1.0.docx"
    if HAS_DOCXTPL and tpl_path.exists():
        logger.info("Using docxtpl template: %s", tpl_path)
        tpl = DocxTemplate(str(tpl_path))
        context["today"] = datetime.now().strftime("%B %d, %Y")
        context["compliance_table"] = context.get("unified_compliance_matrix", [])
        tpl.render(context)
        tpl.save(str(out_path))
    else:
        logger.info("Building styled DOCX from scratch...")
        doc = _build_from_scratch(context, visuals_dir)
        doc.save(str(out_path))

    logger.info("Rich DOCX exported -> %s", out_path)
    return out_path
